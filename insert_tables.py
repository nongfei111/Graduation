# -*- coding: utf-8 -*-
"""
Script to insert all 10 tables into the Chinese translation document.
"""
import docx
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy
import os

SOURCE = r"C:\Users\HP\Desktop\论文材料\2200330230农飞外文翻译译文_标准格式.docx"
OUTPUT = r"C:\Users\HP\Desktop\论文材料\2200330230农飞外文翻译译文_完整版.docx"

# ===================== TABLE DATA DEFINITIONS =====================

# Table 1: Application domains for facial biometric detection
TABLE1_TITLE = "表1. 人脸生物识别检测的几个应用领域"
TABLE1_HEADERS = ["应用领域", "用途"]
TABLE1_DATA = [
    ["安防与执法", "识别和追踪犯罪分子，加速调查取证"],
    ["银行与零售", "通过eKYC进行客户验证，用于银行应用认证、无卡ATM交易、在线账户创建以及Apple Pay等数字支付[7,8]"],
    ["医疗健康领域", "检测遗传疾病，追踪患者用药效果，健康保险记录管理[9]"],
    ["移民与边境检查", "在欧洲各国机场实施人脸识别身份核验[10]"],
    ["教育", "校园安全、考勤监控、提高学习参与度[11]"],
    ["移动设备", "智能手机中的FaceID，如Apple、Samsung、Motorola和OnePlus[12]"],
]

# Table 2: Prior research related to FLD
TABLE2_TITLE = "表2. 与FLD相关的前人研究"
TABLE2_HEADERS = ["参考文献", "目标与主题", "观察与局限性", "类型"]
TABLE2_DATA = [
    ["[28] 2017",
     "包含人脸呈现攻击检测的最新方法及该领域各实验室的概述，同时描述了该领域的挑战与竞赛",
     "未遵循PRISMA方法，关注2011年至2017年的挑战与竞赛",
     "综述"],
    ["[27] 2019",
     "讨论了使用PRISMA方法的系统综述，侧重活体指示线索，特别是作为确定各类欺骗问题最佳解决方案的指南",
     "仅综述了2014-2017年间发表的65篇文章，仅关注活体指示线索，缺乏2017年以来该领域的新趋势",
     "系统文献综述"],
    ["[24] 2020",
     "讨论了呈现攻击的类型学及2D和3D攻击的各种可用数据库中的检测方法，涵盖PAD面临的挑战、演变和当前趋势，并提供了未来研究的新视角",
     "未遵循PRISMA方法，更多关注基于RGB的方法，基于传感器的PAD方法未详细探讨，需要探索更先进的研究方向",
     "综述"],
    ["[25] 2020",
     "讨论了纹理、运动、基于多融合的人脸反欺骗方法及可用的2D攻击数据库，描述了包括CNN、纹理特征描述符和基于运动的技术在内的各种人脸反欺骗技术",
     "未遵循PRISMA方法，仅讨论2015-2019年最近四年，仅关注基于深度学习的解决方案，数据库未广泛综述，未来方向未讨论",
     "综述"],
    ["[26] 2020",
     "包含基于图像纹理、图像质量、计算机交互、深度分析的显式特征提取方法，以及深度学习特征提取、迁移学习、特征集成和领域泛化",
     "研究中使用了少量研究文章，缺乏对PAD可用数据库的讨论，探索的未来方向非常有限",
     "综述"],
    ["[36] 2021",
     "讨论了单模态和多模态人脸呈现攻击检测的国际竞赛",
     "包含2019年至2021年的最新五次竞赛，未遵循PRISMA方法",
     "综述"],
    ["[30] 2021",
     "包含用于FPAD的先进深度学习和多模态融合方法，进行了深入的技术审查，包括最新的深度学习方法、数据集和评估指标",
     "未遵循PRISMA方法，侧重深入的深度学习和多模态方法，包括截至2021年的研究文章",
     "综述"],
    ["[31] 2021",
     "包含用于人脸反欺骗问题的深度学习方法及数据集，讨论了基于传感器方法的技术",
     "未遵循PRISMA方法，更多关注深度学习方法，需要探索先进的研究方向",
     "综述"],
]

# Table 3: Research Goals
TABLE3_TITLE = "表3. 研究目标"
TABLE3_HEADERS = ["编号", "研究问题", "动机", "回答章节"]
TABLE3_DATA = [
    ["RQ1", "与人脸活体检测方法相关的已发表论文按年份、出版物和出版物类型的分布情况如何？",
     "有助于确定研究在何时、何地以及由谁开展", "第3节，第3.1节"],
    ["RQ2", "针对人脸识别系统的各种攻击有哪些？\n有哪些不同类型呈现攻击的可用数据集？",
     "有助于探索对人脸识别系统实施的不同类型攻击；有助于定位具有适当训练和测试数据的数据集以取得良好研究成果",
     "第3节，第3.2节"],
    ["RQ3", "有哪些不同类型呈现攻击的可用数据集？",
     "有助于定位具有适当训练和测试数据的数据集以取得良好研究成果",
     "第3节，第3.3节"],
    ["RQ4", "与人脸呈现攻击检测相关的人工智能主要方法有哪些？人脸活体检测中使用的评估指标是什么？",
     "有助于识别适用于当今人脸生物识别应用的适当人工智能方法；有助于选择合适的评估指标进行性能度量",
     "第4节，第5节"],
    ["RQ5", "现有人脸反欺骗技术面临的主要挑战和问题是什么？",
     "有助于探索研究人脸呈现攻击时出现的根本问题，以及当前解决方案的优势和局限性",
     "第6节，第6.1节，第6.2节，第7节"],
    ["RQ6", "鲁棒且可靠的人脸活体检测系统的未来方向是什么？",
     "有助于发现尚未探索的重要研究方向",
     "第7节"],
]

# Table 4: Primary and secondary keywords
TABLE4_TITLE = "表4. 主要和次要关键词列表"
TABLE4_HEADERS = ["关键词类别", "关键词"]
TABLE4_DATA = [
    ["基础关键词", "\"Face Anti-Spoofing\""],
    ["主要关键词", "\"Face Liveness detection\", \"Face Presentation Attacks\", \"Face Anti-spoofing\""],
    ["次要关键词", "\"Artificial Intelligence\", \"Domain Adaptation\", \"Machine Learning\", \"Deep Learning\", \"Domain Generalization\", \"Reinforcement Learning\", \"Face Biometric spoofing\""],
]

# Table 5: Search queries
TABLE5_TITLE = "表5. 搜索查询"
TABLE5_HEADERS = ["数据库", "查询语句", "初始结果数"]
TABLE5_DATA = [
    ["Scopus", "((machine learning) OR (Deep Learning) OR (Artificial Intelligence) OR (Domain Adaptation) OR (Domain Generalization) OR (Reinforcement Learning)) AND ((Face Anti Spoofing) OR (Face Presentation Attacks) OR (Face Liveness Detection) OR (Face Biometric Spoofing))", "283"],
    ["Web of Science (WoS)", "((machine learning) OR (Deep Learning) OR (Artificial Intelligence) OR (Domain Adaptation) OR (Domain Generalization) OR (Reinforcement Learning)) AND ((Face Anti Spoofing) OR (Face Presentation Attacks) OR (Face Liveness Detection) OR (Face Biometric Spoofing))", "188"],
]

# Table 6: Inclusion and exclusion criteria
TABLE6_TITLE = "表6. 纳入和排除标准汇总"
TABLE6_HEADERS = ["纳入标准", "排除标准"]
TABLE6_DATA = [
    ["文章应为原创研究文章，而非综述/调查类文章", "非英文撰写的文章"],
    ["研究文章应在2011年至2022年期间发表", "重复的研究文章"],
    ["研究论文/文章应在标题、摘要或全文中包含搜索关键词", "全文不可获取的研究文章"],
    ["研究文章应至少回答一个研究问题", "与人脸活体检测、人脸呈现攻击、人脸反欺骗无关的研究文章"],
    ["所开发的解决方案应旨在解决人脸呈现攻击检测问题", ""],
]

# Table 7: 2D & 3D attack Datasets (simplified - the most important ones)
TABLE7_TITLE = "表7. 研究中使用的公开可用的2D和3D攻击数据集"
TABLE7_HEADERS = ["数据集", "年份", "主体数", "样本数\n(真实/伪造)", "攻击类型与模式\n(静态/动态, 2D/3D)", "分辨率", "来源机构", "文献引用"]
TABLE7_DATA = [
    ["NUAA [43]", "2008", "15", "5105/7509", "照片攻击 (2D静态)", "640×480", "南京航空航天大学", "[44–52]"],
    ["Replay-Attack [40]", "2011-2012", "50", "300/1000", "照片攻击 (2D静态)/视频重放攻击 (2D动态)", "320×240", "IDIAP研究所", "[53–64]"],
    ["CASIA-FASD [56]", "2012", "50", "150/450", "照片攻击 (切割/打印/卷曲)/视频重放攻击", "多种分辨率", "中国科学院自动化研究所", "[64–68]"],
    ["MSU-MFSD [74]", "2015", "35", "110/330", "打印照片攻击/视频攻击 (2D)", "1920×1080", "密歇根州立大学", "[60,75–81]"],
    ["OULU-NPU [39]", "2017", "55", "990/3960v", "照片攻击/视频重放攻击 (2D)", "多种分辨率", "奥卢大学", "[90]"],
    ["SiW [95]", "2018", "165", "1320/3300v", "打印纸张攻击 (2D)", "1920×1080", "密歇根州立大学", "[71]"],
    ["CASIA-SURF [96]", "2019", "1000", "3000/18000v", "平面切割/卷曲切割照片 (2D)", "RealSense RGB 1280×720", "中国科学院", "[76,77,80]"],
    ["WMCA [78]", "2019", "72", "1679(347/1332)", "2D打印/视频重放/面具攻击 (2D, 3D)", "多种分辨率 1260×720, 320×240", "Idiap研究所", "[79,97]"],
    ["CASIA-SURF CeFA [98]", "2020", "1607 (3个民族子集)", "1800/5400v", "打印/重放/3D打印/红外攻击 (2D & 3D)", "299×299", "中国科学院", "[82]"],
    ["HiFiMask [100]", "2021", "75", "54600v", "3D面具攻击", "—", "中国科学院", "[84]"],
    ["ROSE-YOUTU [73]", "2018", "20", "3350", "打印纸张攻击/视频重放攻击/纸张面具攻击/裁剪面具/全面具/上面具", "640×480", "腾讯公司与NTU ROSE实验室", "[74]"],
    ["SURF [96]", "—", "—", "—", "平面切割/卷曲切割", "—", "中国科学院自动化研究所", "[96]"],
]

# Table 8: Categorization of feature extraction methods
TABLE8_TITLE = "表8. 人脸活体检测特征提取方法分类"
TABLE8_HEADERS = ["提取的特征", "使用的方法", "识别的攻击类型"]
TABLE8_DATA = [
    ["基于纹理的特征提取\n(从输入图像)", "手工纹理特征：LBP[102]、Color LBP[112]、LBP+Gabor Wavelets、LBP+GW+HoG[112]\n频率纹理：2D FFT\n空间纹理：DoG[111]\nCNN方法：AlexNet[118–120]、Fine Tune VGG-Face[121]、DPCNN、FaceDs[102]、DeepPixBis[122]",
     "照片攻击、视频重放攻击、低质量3D面具攻击"],
    ["基于动态纹理", "时空纹理[120]\n动态模式分解(DMD)[86]\nCNN方法：LSTM-CNN[89]、STASN[121]",
     "照片攻击、视频重放攻击"],
    ["基于非侵入式运动特征提取", "眨眼检测[58,86,116]\n头部运动[94]\n频率动态分配器(FDD)[96]\n光流线(OFL)[122]",
     "静态照片攻击"],
    ["基于侵入式运动特征提取", "唇语识别(OFL)[123]",
     "照片攻击与3D面具攻击、低质量视频重放攻击"],
    ["基于rPPG运动特征提取", "rPPG频谱[124]\n局部rPPG相关性[125]\nDeep rPPG[125]",
     "照片攻击、视频重放攻击"],
    ["基于3D形状(深度)的特征提取", "稀疏3D人脸重建[117]\nCNN方法、NAS方法[109]、3D点云网络[126]",
     "平面照片攻击、视频重放攻击"],
]

# Table 9: Popular methodologies in FLD (abbreviated)
TABLE9_TITLE = "表9. 人脸活体检测中常用的一些方法"
TABLE9_HEADERS = ["参考文献与年份", "方法", "DA/DG", "数据集", "性能指标与模型性能", "数据库内测试", "跨数据库测试"]
TABLE9_DATA = [
    ["[58] 2015", "个性化领域自适应", "DA", "CASIA、Replay-Attack", "HTER: 1.40% (Replay-Attack), 10.54% (CASIA)", "Y", "N"],
    ["[97] 2018", "无监督领域自适应框架", "DA", "自有数据集ROSE-YOUTU", "HTER: 27.70%", "Y", "N"],
    ["[76] 2019", "双重三元组挖掘约束下的泛化(OCA-FAS)", "DA", "Replay-Attack、MSU-MFSD、Oulu-NPU", "HTER: 27.98%, AUC: 80.02%", "Y", "Y"],
    ["[133] 2020", "全卷积网络与领域自适应和尺度自适应(FCN-DA-LSA)", "DA", "CASIA-FASD、Replay-Attack、OULU-NPU", "ACER: 1.69%", "Y", "Y"],
    ["[134] 2020", "使用领域引导剪枝的单类领域自适应CNN", "DA", "OULU-NPU、Replay-Mobile、WMCA、IJB-C", "HTER: 21.83%", "Y", "Y"],
    ["[135] 2020", "单边领域泛化框架(SSDG)", "DG", "OULU-NPU、CASIA-FASD、Replay-Attack、MSU-MFSD", "AUC, ROC, APCER", "Y", "Y"],
    ["[136] 2020", "领域无关特征学习", "DG", "Oulu-NPU、CASIA-MFSD、Replay-Attack、MSU-MFSD", "HTER: 7.38%, AUC: 97.17%", "Y", "Y"],
    ["[137] 2020", "全对混淆(TPC)损失和快速领域自适应(FDA)", "DG", "CASIA-FASD、Replay-Attack、MSU-MFSD、Oulu-NPU、SiW", "HTER: 14.00%, ACER: 8.05%", "N", "Y"],
    ["[122] 2020", "无监督对抗领域自适应与解纠缠表示(DR-UDA)", "DA", "Replay-Attack、CASIA、MSU-MFSD、Oulu-NPU、SiW", "HTER: 26.30%", "Y", "Y"],
    ["[139] 2021", "使用解纠缠表示的无监督对抗领域自适应", "DA", "MSU-MFSD、ROSE-Youtu、Oulu-NPU (CASIA-SURF RGB)", "HTER: 28.70%, EER: 3.20%", "Y", "Y"],
    ["[53] 2021", "自监督区域全卷积网络(SSR-FCN)", "DG", "SiW-M、Oulu-NPU、CASIA-FASD、Replay-Attack", "ACER: 2.80%, HTER: 19.90%", "Y", "Y"],
    ["[140] 2021", "相机不变特征学习用于泛化人脸反欺骗", "DG", "CASIA-FASD、Replay-Attack、Oulu-NPU、MSU-MFSD", "EER: 0.89%, HTER: 17.60%", "Y", "Y"],
]

# Table 10: Anomaly detection methods in FLD
TABLE10_TITLE = "表10. FLD中的异常检测方法"
TABLE10_HEADERS = ["参考文献与年份", "异常检测方法", "数据集", "性能指标与模型性能", "数据库内测试", "跨数据库测试"]
TABLE10_DATA = [
    ["[141] 2018", "GMM异常检测器与聚合数据库", "Replay-Mobile、Replay-Attack、OULU-NPU、MSU-MFSD", "—", "Y", "Y"],
    ["[140] 2019", "深度度量学习模型", "Replay-Attack、Replay-Mobile、MSU-MFSD", "—", "Y", "Y"],
    ["[142] 2020", "超球面损失函数", "Replay-Mobile、Replay-Attack、OULU-NPU、MSU-MFSD", "—", "Y", "Y"],
    ["[143] 2020", "HOG人脸检测+VGG Face特征提取+伪负样本采样", "Replay-Attack、Replay-Mobile", "—", "Y", "Y"],
    ["[109] 2021", "多核融合用于不可见呈现攻击的异常检测", "—", "—", "Y", "Y"],
    ["[144] 2021", "客户端特定单类", "Replay-Mobile、Replay-Attack、OULU-NPU、MSU-MFSD", "—", "Y", "Y"],
]

# ===================== HELPER FUNCTIONS =====================

def set_cell_font(cell, text, font_name='宋体', font_size=Pt(10.5), bold=False):
    """Set cell text with proper font formatting."""
    # Clear existing paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    
    # Use first paragraph
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)
    
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold

def set_cell_font_left(cell, text, font_name='宋体', font_size=Pt(10.5), bold=False):
    """Set cell text with left alignment."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(14)
    
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.font.bold = bold

def create_table_data(doc, headers, data, table_title_paragraph=None):
    """Create a docx table and return it."""
    rows = len(data) + 1  # +1 for header
    cols = len(headers)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table font for all cells
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            if row_idx == 0:
                # Header row
                set_cell_font(cell, headers[col_idx], bold=True)
                # Add shading to header
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                cell._element.get_or_add_tcPr().append(shading_elm)
            else:
                if col_idx < len(data[row_idx - 1]):
                    text = data[row_idx - 1][col_idx]
                else:
                    text = ""
                # If text contains newlines, use left alignment
                if '\n' in text:
                    set_cell_font_left(cell, text, bold=False)
                else:
                    set_cell_font(cell, text, bold=False)
    
    return table

def insert_table_after_paragraph(paragraph, table):
    """Insert a table after a given paragraph using XML manipulation."""
    # Add the table after the paragraph in the document body
    paragraph._element.addnext(table._tbl)

def delete_paragraphs_between(doc, start_idx, end_idx):
    """Delete paragraphs between start_idx and end_idx (inclusive)."""
    # Get the body element
    body = doc.element.body
    paras_to_remove = []
    for i in range(start_idx, end_idx + 1):
        if i < len(doc.paragraphs):
            p_elem = doc.paragraphs[i]._element
            paras_to_remove.append(p_elem)
    for elem in paras_to_remove:
        body.remove(elem)

# ===================== MAIN SCRIPT =====================

def main():
    print("Loading translation document...")
    doc = Document(SOURCE)
    
    # ===== TABLE 1: Application Domains (ALREADY EXISTS as a real table in the doc) =====
    print("Processing Table 1: Application Domains... (already exists, skipping)")
    # Table 1 is already present as a real Word table in the translation document
    # No action needed - it has proper data rows
    
    # ===== TABLE 2: Prior Research =====
    print("Processing Table 2: Prior Research...")
    t2_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == TABLE2_TITLE:
            t2_idx = i
            break
    
    if t2_idx is not None:
        # Delete any placeholder paragraphs after Table 2 title (old text remnants like "Ref. Objectives...")
        body = doc.element.body
        paras_to_delete = []
        for j in range(t2_idx + 1, min(t2_idx + 20, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip()
            if txt == '' or 'Ref.' in txt[:10] or 'Objectives' in txt[:15] or 'Observation' in txt[:15] or \
               txt in ['Review', 'Systematic Literature Review'] or \
               (txt.startswith('[') and len(txt) < 15 and ']' in txt[:6]):
                paras_to_delete.append(j)
        
        for j in reversed(paras_to_delete):
            body.remove(doc.paragraphs[j]._element)
        
        # Refresh
        t2_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == TABLE2_TITLE:
                t2_idx = i
                break
        
        if t2_idx is not None:
            table2 = create_table_data(doc, TABLE2_HEADERS, TABLE2_DATA)
            insert_table_after_paragraph(doc.paragraphs[t2_idx], table2)
            print("  Table 2 inserted successfully.")
    else:
        print("  WARNING: Table 2 title not found!")
    
    # ===== TABLE 3: Research Goals =====
    print("Processing Table 3: Research Goals...")
    t3_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == TABLE3_TITLE:
            t3_idx = i
            break
    
    if t3_idx is not None:
        body = doc.element.body
        paras_to_delete = []
        for j in range(t3_idx + 1, min(t3_idx + 15, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip()
            if txt == '' or 'Number' in txt[:10] or 'Research Questions' in txt[:20] or \
               txt.startswith('RQ') or 'Motivations' in txt[:15] or 'Answered' in txt[:15] or \
               txt in ['Section 3', 'Section 4', 'Section 5', 'Section 6'] or \
               'Section 3,' in txt[:15] or 'Section 4,' in txt[:15] or 'Section 6,' in txt[:15] or \
               'Section 6.2' in txt[:15] or 'Section 7' in txt[:15]:
                paras_to_delete.append(j)
        
        for j in reversed(paras_to_delete):
            body.remove(doc.paragraphs[j]._element)
        
        t3_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == TABLE3_TITLE:
                t3_idx = i
                break
        
        if t3_idx is not None:
            table3 = create_table_data(doc, TABLE3_HEADERS, TABLE3_DATA)
            insert_table_after_paragraph(doc.paragraphs[t3_idx], table3)
            print("  Table 3 inserted successfully.")
    else:
        print("  WARNING: Table 3 title not found!")
    
    # ===== TABLE 4 & TABLE 5: Keywords and Search Queries (both referenced in same paragraph P67) =====
    print("Processing Table 4: Keywords...")
    t45_ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '表4列出了基本、主要和次要关键词' in p.text:
            t45_ref_idx = i
            break
    
    if t45_ref_idx is not None:
        from docx.oxml import OxmlElement
        
        # Create Table 4 title paragraph
        new_p4 = OxmlElement('w:p')
        new_p4_pr = OxmlElement('w:pPr')
        new_p4.append(new_p4_pr)
        new_r4 = OxmlElement('w:r')
        new_r4Pr = OxmlElement('w:rPr')
        new_r4.append(new_r4Pr)
        new_t4 = OxmlElement('w:t')
        new_t4.text = TABLE4_TITLE
        new_t4.set(qn('xml:space'), 'preserve')
        new_r4.append(new_t4)
        new_p4.append(new_r4)
        
        # Create Table 5 title paragraph
        new_p5 = OxmlElement('w:p')
        new_p5_pr = OxmlElement('w:pPr')
        new_p5.append(new_p5_pr)
        new_r5 = OxmlElement('w:r')
        new_r5Pr = OxmlElement('w:rPr')
        new_r5.append(new_r5Pr)
        new_t5 = OxmlElement('w:t')
        new_t5.text = TABLE5_TITLE
        new_t5.set(qn('xml:space'), 'preserve')
        new_r5.append(new_t5)
        new_p5.append(new_r5)
        
        # Insert Table 4 title after reference paragraph, then Table 5 title after Table 4 title
        # Order: ref_para -> Table4_title -> Table4_table -> Table5_title -> Table5_table
        doc.paragraphs[t45_ref_idx]._element.addnext(new_p4)
        
        # Create and insert Table 4
        table4 = doc.add_table(rows=len(TABLE4_DATA)+1, cols=len(TABLE4_HEADERS))
        table4.style = 'Table Grid'
        table4.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table4.rows):
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    set_cell_font(cell, TABLE4_HEADERS[col_idx], bold=True)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    set_cell_font_left(cell, TABLE4_DATA[row_idx-1][col_idx])
        new_p4.addnext(table4._tbl)
        
        # Insert Table 5 title after Table 4
        table4._tbl.addnext(new_p5)
        
        # Create and insert Table 5
        table5 = doc.add_table(rows=len(TABLE5_DATA)+1, cols=len(TABLE5_HEADERS))
        table5.style = 'Table Grid'
        table5.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table5.rows):
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    set_cell_font(cell, TABLE5_HEADERS[col_idx], bold=True)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    set_cell_font_left(cell, TABLE5_DATA[row_idx-1][col_idx])
        new_p5.addnext(table5._tbl)
        
        print("  Table 4 and Table 5 inserted successfully.")
    else:
        print("  WARNING: Table 4/5 reference not found!")
    
    # ===== TABLE 6: Inclusion and Exclusion Criteria =====
    print("Processing Table 6: Inclusion/Exclusion Criteria...")
    t6_ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '如表6所示，作者为研究论文选择建立了一套纳入标准' in p.text:
            t6_ref_idx = i
            break
    
    if t6_ref_idx is not None:
        from docx.oxml import OxmlElement
        new_p = OxmlElement('w:p')
        new_p_pr = OxmlElement('w:pPr')
        new_p.append(new_p_pr)
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.text = TABLE6_TITLE
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        
        doc.paragraphs[t6_ref_idx]._element.addnext(new_p)
        
        table6 = doc.add_table(rows=len(TABLE6_DATA)+1, cols=len(TABLE6_HEADERS))
        table6.style = 'Table Grid'
        table6.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(table6.rows):
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    set_cell_font(cell, TABLE6_HEADERS[col_idx], bold=True)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    set_cell_font_left(cell, TABLE6_DATA[row_idx-1][col_idx])
        
        new_p.addnext(table6._tbl)
        print("  Table 6 inserted successfully.")
    else:
        print("  WARNING: Table 6 reference not found!")
    
    # ===== TABLE 7: 2D & 3D Attack Datasets =====
    print("Processing Table 7: 2D & 3D Attack Datasets...")
    t7_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == TABLE7_TITLE:
            t7_idx = i
            break
    
    if t7_idx is not None:
        # Delete old text-format table data (P103-P112)
        body = doc.element.body
        paras_to_delete = []
        for j in range(t7_idx + 1, min(t7_idx + 15, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip()
            # These are the dataset description lines
            if txt and ('(' in txt[:3] or 'NUAA' in txt[:10] or 'Replay-Attack' in txt[:20] or 
                       'CASIA-FASD' in txt[:20] or 'MSU-MFSD' in txt[:20] or 
                       'OULU-NPU' in txt[:20] or 'SiW' in txt[:10] or 'CASIA-SURF' in txt[:20] or
                       'WMCA' in txt[:10] or 'HiFiMask' in txt[:20] or 'ROSE' in txt[:10]):
                paras_to_delete.append(j)
        
        for j in reversed(paras_to_delete):
            body.remove(doc.paragraphs[j]._element)
        
        # Refresh
        t7_idx = None
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == TABLE7_TITLE:
                t7_idx = i
                break
        
        if t7_idx is not None:
            table7 = create_table_data(doc, TABLE7_HEADERS, TABLE7_DATA)
            insert_table_after_paragraph(doc.paragraphs[t7_idx], table7)
            print("  Table 7 inserted successfully.")
    else:
        print("  WARNING: Table 7 title not found!")
    
    # ===== TABLE 8: Feature Extraction Methods =====
    print("Processing Table 8: Feature Extraction Methods...")
    # Find the paragraph just before where Table 8 should be
    # In original, it's after "image quality-based" section, around P133 (translation)
    # Find section 4.1.5 "现有技术中的问题" and insert before it
    t8_ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '4.1.5 现有技术中的问题' in p.text:
            t8_ref_idx = i
            break
    
    if t8_ref_idx is not None:
        from docx.oxml import OxmlElement
        new_p = OxmlElement('w:p')
        new_p_pr = OxmlElement('w:pPr')
        new_p.append(new_p_pr)
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.text = TABLE8_TITLE
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        
        # Insert before the 4.1.5 section
        doc.paragraphs[t8_ref_idx]._element.addprevious(new_p)
        
        table8 = doc.add_table(rows=len(TABLE8_DATA)+1, cols=len(TABLE8_HEADERS))
        table8.style = 'Table Grid'
        table8.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(table8.rows):
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    set_cell_font(cell, TABLE8_HEADERS[col_idx], bold=True)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    text = TABLE8_DATA[row_idx-1][col_idx]
                    if '\n' in text:
                        set_cell_font_left(cell, text)
                    else:
                        set_cell_font(cell, text)
        
        new_p.addnext(table8._tbl)
        print("  Table 8 inserted successfully.")
    else:
        print("  WARNING: Table 8 insertion point not found!")
    
    # ===== TABLE 9: Popular Methodologies =====
    print("Processing Table 9: Popular Methodologies...")
    t9_ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '表9总结了人脸活体检测的最新研究' in p.text:
            t9_ref_idx = i
            break
    
    if t9_ref_idx is not None:
        from docx.oxml import OxmlElement
        new_p = OxmlElement('w:p')
        new_p_pr = OxmlElement('w:pPr')
        new_p.append(new_p_pr)
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.text = TABLE9_TITLE
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        
        doc.paragraphs[t9_ref_idx]._element.addnext(new_p)
        
        table9 = doc.add_table(rows=len(TABLE9_DATA)+1, cols=len(TABLE9_HEADERS))
        table9.style = 'Table Grid'
        table9.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(table9.rows):
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    set_cell_font(cell, TABLE9_HEADERS[col_idx], bold=True)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    set_cell_font(cell, TABLE9_DATA[row_idx-1][col_idx])
        
        new_p.addnext(table9._tbl)
        print("  Table 9 inserted successfully.")
    else:
        print("  WARNING: Table 9 reference not found!")
    
    # ===== TABLE 10: Anomaly Detection =====
    print("Processing Table 10: Anomaly Detection...")
    t10_ref_idx = None
    for i, p in enumerate(doc.paragraphs):
        if '表10总结了FLD中的异常检测方法' in p.text:
            t10_ref_idx = i
            break
    
    if t10_ref_idx is not None:
        from docx.oxml import OxmlElement
        new_p = OxmlElement('w:p')
        new_p_pr = OxmlElement('w:pPr')
        new_p.append(new_p_pr)
        new_r = OxmlElement('w:r')
        new_rPr = OxmlElement('w:rPr')
        new_r.append(new_rPr)
        new_t = OxmlElement('w:t')
        new_t.text = TABLE10_TITLE
        new_t.set(qn('xml:space'), 'preserve')
        new_r.append(new_t)
        new_p.append(new_r)
        
        doc.paragraphs[t10_ref_idx]._element.addnext(new_p)
        
        table10 = doc.add_table(rows=len(TABLE10_DATA)+1, cols=len(TABLE10_HEADERS))
        table10.style = 'Table Grid'
        table10.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for row_idx, row in enumerate(table10.rows):
            for col_idx, cell in enumerate(row.cells):
                if row_idx == 0:
                    set_cell_font(cell, TABLE10_HEADERS[col_idx], bold=True)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    set_cell_font(cell, TABLE10_DATA[row_idx-1][col_idx])
        
        new_p.addnext(table10._tbl)
        print("  Table 10 inserted successfully.")
    else:
        print("  WARNING: Table 10 reference not found!")
    
    # Save the document
    print(f"\nSaving document to: {OUTPUT}")
    doc.save(OUTPUT)
    print("Done! All tables have been inserted.")

if __name__ == '__main__':
    main()
