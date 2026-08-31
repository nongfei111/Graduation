# -*- coding: utf-8 -*-
"""
Part 1: Setup, page config, cover page, paper info, Introduction section
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ============================================================
# PAGE SETUP: A4, margins 2.5cm all sides
# ============================================================
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Set default paragraph format
pf = style.paragraph_format
pf.line_spacing = Pt(20)  # 固定值20磅

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def set_font(run, font_name_cn='宋体', font_name_en='Times New Roman', size=Pt(12), bold=False):
    """Set font for a run with both Chinese and English fonts"""
    run.font.size = size
    run.bold = bold
    run.font.name = font_name_en
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_cn)
    rFonts.set(qn('w:ascii'), font_name_en)
    rFonts.set(qn('w:hAnsi'), font_name_en)
    run.font.name = font_name_en

def add_paragraph_with_format(doc, text, font_cn='宋体', font_en='Times New Roman', 
                                size=Pt(12), bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                first_line_indent=Cm(0.74), space_before=Pt(0), space_after=Pt(0),
                                line_spacing=Pt(20)):
    """Add a paragraph with proper formatting"""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = first_line_indent
    run = p.add_run(text)
    set_font(run, font_cn, font_en, size, bold)
    return p

def add_heading_0(doc, text):
    """0级标题: 黑体 四号(14pt), centered, space before/after 1 line"""
    return add_paragraph_with_format(doc, text, '黑体', 'Times New Roman', 
                                     Pt(14), True, WD_ALIGN_PARAGRAPH.CENTER,
                                     None, Pt(20), Pt(20), Pt(20))

def add_heading_1(doc, text):
    """1级标题: 黑体 四号(14pt), bold, left-aligned"""
    return add_paragraph_with_format(doc, text, '黑体', 'Times New Roman',
                                     Pt(14), True, WD_ALIGN_PARAGRAPH.LEFT,
                                     Cm(0), Pt(12), Pt(6), Pt(20))

def add_heading_2(doc, text):
    """2级标题: 黑体 小四号(12pt), bold"""
    return add_paragraph_with_format(doc, text, '黑体', 'Times New Roman',
                                     Pt(12), True, WD_ALIGN_PARAGRAPH.LEFT,
                                     Cm(0), Pt(8), Pt(4), Pt(20))

def add_heading_3(doc, text):
    """3级标题: 楷体 小四号(12pt)"""
    return add_paragraph_with_format(doc, text, '楷体', 'Times New Roman',
                                     Pt(12), True, WD_ALIGN_PARAGRAPH.LEFT,
                                     Cm(0), Pt(6), Pt(3), Pt(20))

def add_body(doc, text):
    """正文: 宋体 小四号(12pt), 首行缩进2字符"""
    return add_paragraph_with_format(doc, text, '宋体', 'Times New Roman',
                                     Pt(12), False, WD_ALIGN_PARAGRAPH.JUSTIFY,
                                     Cm(0.74), Pt(0), Pt(0), Pt(20))

def add_ref_entry(doc, text):
    """参考文献条目: 宋体 五号(10.5pt)"""
    return add_paragraph_with_format(doc, text, '宋体', 'Times New Roman',
                                     Pt(10.5), False, WD_ALIGN_PARAGRAPH.JUSTIFY,
                                     Cm(0), Pt(0), Pt(0), Pt(20))

def add_table_title(doc, text):
    """表格标题: 黑体 五号(10.5pt), 居中"""
    return add_paragraph_with_format(doc, text, '黑体', 'Times New Roman',
                                     Pt(10.5), True, WD_ALIGN_PARAGRAPH.CENTER,
                                     None, Pt(6), Pt(3), Pt(20))

def add_fig_label(doc, text):
    """图注"""
    return add_paragraph_with_format(doc, text, '宋体', 'Times New Roman',
                                     Pt(10.5), False, WD_ALIGN_PARAGRAPH.CENTER,
                                     None, Pt(3), Pt(3), Pt(20))

def set_cell_font(cell, text, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5), bold=False):
    """Set cell content with proper font"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = Pt(15)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, font_cn, font_en, size, bold)

def create_table_title_row(table, headers, cols):
    """Create table header row"""
    from docx.shared import Pt
    for i, header in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], header, '黑体', 'Times New Roman', Pt(10.5), True)

# ============================================================
# COVER PAGE (封面)
# ============================================================
# Add empty paragraphs for centering (approximate vertical centering)
for _ in range(6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(20)

# Title: 外文翻译译文 - 黑体 三号(16pt)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = Pt(20)
run = p.add_run('外文翻译译文')
set_font(run, '黑体', 'Times New Roman', Pt(16), True)

# English title - Times New Roman 小四号(12pt)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.line_spacing = Pt(20)
p.paragraph_format.space_before = Pt(12)
run = p.add_run('Face Liveness Detection Using Artificial Intelligence Techniques:\nA Systematic Literature Review and Future Directions')
set_font(run, 'Times New Roman', 'Times New Roman', Pt(12), False)

# Blank line
p = doc.add_paragraph()
p.paragraph_format.line_spacing = Pt(20)

# Student info
info_lines = [
    ('学    号：', '2200330230'),
    ('姓    名：', '农飞'),
    ('学    院：', '计算机与信息学院'),
    ('专    业：', '物联网工程'),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = Pt(30)
    run1 = p.add_run(label)
    set_font(run1, '宋体', 'Times New Roman', Pt(14), False)
    run2 = p.add_run(value)
    set_font(run2, '宋体', 'Times New Roman', Pt(14), False)

print("Part 1: Cover page created.")
doc.save(r'C:\Users\HP\Desktop\论文材料\temp_part1.docx')
print("Part 1 saved.")
